from openpyxl import Workbook
from docx import Document

# Create Excel workbook
wb = Workbook()
ws = wb.active
ws.title = "Daily Task Sheet"
headers = ["Date", "Team Member", "Task Description", "Start Time", "End Time", "Status", "Remarks"]
ws.append(headers)
wb.save("C:/Users/shubhamamadane/Documents/Daily_Task_Sheet.xlsx")

# Create Word document
doc = Document()
doc.add_heading('Daily Task Sheet', 0)
doc.add_paragraph('Date: [Insert Date]')
doc.add_paragraph('Team: [Insert Team Name]')
doc.add_paragraph('')

# Add table to Word doc
table = doc.add_table(rows=1, cols=len(headers))
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

doc.save("C:/Users/shubhamamadane/Documents/Daily_Task_Sheet.docx")

print("Files created: Daily_Task_Sheet.xlsx and Daily_Task_Sheet.docx")
